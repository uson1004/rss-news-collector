from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("final-project-report.docx")

NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF4FA"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "666666"
DARK = "222222"
WHITE = "FFFFFF"
LINE = "B4C6E7"

BODY_FONT = "맑은 고딕"
MONO_FONT = "D2Coding"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border(paragraph, color=LINE, size=8, space=5, side="bottom") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_box(paragraph, color=LINE, size=6, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), str(space))
        border.set(qn("w:color"), color)
        p_bdr.append(border)


def set_repeatable_style_font(style, name: str, size: float, color=DARK, bold=False) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_repeatable_style_font(normal, BODY_FONT, 10.5, DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.6
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = styles["Title"]
    set_repeatable_style_font(title, BODY_FONT, 25, NAVY, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(9)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    set_repeatable_style_font(subtitle, BODY_FONT, 13, MID_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    set_repeatable_style_font(h1, BODY_FONT, 15, NAVY, True)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_repeatable_style_font(h2, BODY_FONT, 12.5, BLUE, True)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_repeatable_style_font(h3, BODY_FONT, 11, NAVY, True)
    h3.paragraph_format.space_before = Pt(5)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        list_style = styles[list_style_name]
        set_repeatable_style_font(list_style, BODY_FONT, 10.5, DARK)
        list_style.paragraph_format.left_indent = Mm(7)
        list_style.paragraph_format.first_line_indent = Mm(-3.5)
        list_style.paragraph_format.space_after = Pt(3)
        list_style.paragraph_format.line_spacing = 1.45


def setup_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("읽을게 | 최종 프로젝트 보고서")
    set_run_font(run, size=8.5, bold=True, color=MID_GRAY)
    set_paragraph_border(p, color="D9E2F3", size=4, space=3)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    add_page_number(first_footer.paragraphs[0])


def add_paragraph(doc, text="", *, bold_lead=None, align=None, after=5, line=1.6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        set_paragraph_border(p, color=LINE, size=8, space=4)
    return p


def add_callout(doc, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(2.5)
    p.paragraph_format.right_indent = Mm(2.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.45
    set_paragraph_shading(p, PALE_BLUE)
    set_paragraph_box(p, color=LINE, size=5, space=5)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=NAVY)
    text_run = p.add_run(text)
    set_run_font(text_run)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent_dxa=0)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=NAVY)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code_block(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(2.5)
    p.paragraph_format.right_indent = Mm(2.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.1
    set_paragraph_shading(p, "F7F7F7")
    set_paragraph_box(p, color="D9D9D9", size=4, space=5)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name=MONO_FONT, size=8.5, color=DARK)
        if idx < len(lines) - 1:
            run.add_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    setup_page(section)
    setup_header_footer(section)

    # Page 1: cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(38)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("FINAL PROJECT REPORT")
    set_run_font(run, size=11, bold=True, color=BLUE)

    doc.add_paragraph("읽을게", style="Title")
    doc.add_paragraph("접근성 중심 RSS 뉴스 리더 및 뉴스레터 서비스", style="Subtitle")

    title_rule = doc.add_paragraph()
    title_rule.paragraph_format.space_after = Pt(22)
    set_paragraph_border(title_rule, color=BLUE, size=12, space=1)

    add_paragraph(
        doc,
        "복잡한 웹페이지와 RSS 콘텐츠를 본문 중심의 읽기 화면으로 재구성하고, "
        "AI 요약·관찰 노트·번역·TTS 및 주간 뉴스레터 기능을 결합한 웹 서비스 개발 결과를 정리하였다.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=25,
        line=1.7,
    )

    meta_rows = [
        ["보고서 구분", "최종 프로젝트 보고서"],
        ["학번", "____________________________"],
        ["이름", "____________________________"],
        ["제출기한", "2026. 6. 21.(일)"],
        ["제출방법", "이메일 제출 (dsm2026@naver.com)"],
        ["소스코드", "이메일 첨부 또는 GitHub 링크 기입: ____________________"],
    ]
    add_table(doc, ["항목", "내용"], meta_rows, [2150, 7210], font_size=9.8)

    add_callout(
        doc,
        "프로젝트 한 줄 요약",
        "URL 및 RSS 기사를 읽기 좋은 화면으로 변환하고, 사용자가 관심 카테고리를 이메일 뉴스레터로 구독할 수 있게 한 서비스",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("React · FastAPI · RSS/Atom · Claude API · DeepL API · SQLite · SMTP")
    set_run_font(run, size=9.5, bold=True, color=MID_GRAY)

    # Page 2: overview and goals
    page_break(doc)
    add_heading(doc, "1. 개요 및 목표", 1)
    add_heading(doc, "1.1 개발 배경", 2)
    add_paragraph(
        doc,
        "일반 웹페이지는 광고, 메뉴, 추천 영역, 팝업 등으로 인해 본문에 집중하기 어렵고, "
        "사이트마다 글자 크기와 줄 간격이 달라 읽기 경험이 일관되지 않다. 특히 긴 기사나 기술 문서를 "
        "읽을 때 사용자가 필요한 것은 원문 자체와 출처, 그리고 자신의 환경에 맞는 읽기 설정이다. "
        "이 문제를 해결하기 위해 URL을 입력하면 핵심 본문을 추출하고, RSS를 통해 읽을 글을 탐색하며, "
        "필요한 경우에만 AI 기능을 사용하는 서비스 ‘읽을게’를 개발하였다.",
    )

    add_heading(doc, "1.2 개발 목표", 2)
    add_number(doc, "URL 또는 RSS 기사에서 제목·출처·본문을 추출해 일관된 리더뷰로 제공한다.")
    add_number(doc, "글자 크기, 줄 간격, 문단 간격, 본문 폭, 화면 테마를 사용자가 조절할 수 있게 한다.")
    add_number(doc, "AI 요약, 관찰 노트, 번역과 TTS를 원문을 보조하는 선택 기능으로 제공한다.")
    add_number(doc, "관심 카테고리를 직접 만들고 이메일 뉴스레터로 구독할 수 있게 한다.")
    add_number(doc, "정기 발송 과정의 중복 처리와 실패 상태를 기록해 운영 가능한 구조를 마련한다.")

    add_heading(doc, "1.3 개발 범위와 판단 기준", 2)
    scope_rows = [
        ["핵심 범위", "URL 본문 추출, RSS 탐색, 공통 리더뷰, 읽기 설정"],
        ["확장 범위", "AI 요약·관찰 노트·번역, TTS, 사용자 카테고리"],
        ["운영 범위", "뉴스레터 구독, 미리보기, SMTP 발송, 주간 배치"],
        ["완료 기준", "주요 API·화면 동작, 백엔드 테스트 통과, 프론트엔드 빌드 성공"],
        ["제외 범위", "사용자 계정, 북마크, 읽기 기록, 전체 Supabase 이전, 실제 배포"],
    ]
    add_table(doc, ["구분", "내용"], scope_rows, [2200, 7160], font_size=9.3)

    add_heading(doc, "1.4 시스템 구성", 2)
    add_paragraph(
        doc,
        "프론트엔드는 React와 Vite로 화면과 상태를 처리하고, 백엔드는 FastAPI가 본문 추출, RSS 파싱, "
        "AI API 호출, 뉴스레터 저장·발송을 담당한다. 서비스 데이터는 SQLite에 저장하며, Supabase는 "
        "공개 테이블 조회를 통한 연결 검증 단계까지만 적용하였다.",
        after=4,
    )
    add_code_block(
        doc,
        [
            "사용자 → React 화면 → FastAPI API",
            "                    ├─ 웹페이지/RSS 본문 추출",
            "                    ├─ Claude·DeepL 외부 API",
            "                    └─ SQLite 저장 → SMTP 발송",
        ],
    )

    # Page 3: functions part 1
    page_break(doc)
    add_heading(doc, "2. 기능 구현 및 결과", 1)
    add_heading(doc, "2.1 URL 본문 추출과 공통 리더뷰", 2)
    add_paragraph(
        doc,
        "사용자가 URL을 입력하면 백엔드가 httpx로 페이지를 요청하고 trafilatura를 우선 사용해 본문을 추출한다. "
        "추출 결과가 부족하면 BeautifulSoup 기반 단락 수집을 적용하며, 사이트가 직접 접근을 차단한 경우에는 "
        "공식 RSS/Atom 피드에서 같은 글을 찾아 본문을 복구한다. 프로토콜을 생략한 주소는 HTTPS로 보정하고 "
        "FTP 등 지원하지 않는 프로토콜은 검증 단계에서 거부한다.",
    )

    feature_rows_1 = [
        ["본문 추출", "제목·출처·설명·본문·단락·어절 수 반환", "구현"],
        ["RSS fallback", "직접 추출 실패 시 피드 항목의 URL·GUID·ID 비교", "구현"],
        ["읽기 설정", "글자/줄/문단/폭과 밝음·어두움·고대비 테마", "구현"],
        ["설정 유지", "브라우저 localStorage에 읽기 설정 저장", "부분 구현"],
        ["오류 안내", "차단·비HTML·RSS 미발견 등 원인별 메시지", "구현"],
    ]
    add_table(doc, ["기능", "구현 결과", "상태"], feature_rows_1, [1800, 5960, 1600], font_size=9.0)

    add_heading(doc, "2.2 RSS 카테고리와 아이디어 레이더", 2)
    add_paragraph(
        doc,
        "국내 뉴스·기술블로그·긱뉴스 등의 RSS를 카테고리별로 수집하고, 키워드 필터와 최신순 정렬을 적용하였다. "
        "뉴스 화면은 첫 기사를 대표 기사로 강조하고 이후 기사를 번호가 있는 목차 형태로 배치하였다. "
        "또한 카테고리 이름과 검색 힌트를 입력하면 내부 피드 목록, 도메인 자동 발견, 필요 시 Claude 기반 힌트 확장을 "
        "통해 RSS 후보를 제안하는 3단계 카테고리 추가 기능을 구현하였다.",
    )

    add_heading(doc, "2.3 AI·번역·음성 기능", 2)
    add_bullet(doc, "AI 요약: 본문의 핵심 내용을 최대 5개의 짧은 항목으로 압축한다.")
    add_bullet(doc, "질문 프롬프트: 생성된 요약을 바탕으로 복습·토론용 질문을 만든다.")
    add_bullet(doc, "AI 관찰 노트: 읽을 이유, 변화 신호, 불편과 욕망, 아이디어 힌트, 주의점을 구조화한다.")
    add_bullet(doc, "번역: DeepL API Free 엔드포인트로 영어·일본어·중국어 번역을 요청한다.")
    add_bullet(doc, "TTS: Web Speech API로 단락을 순서대로 읽고 일시정지·재개·정지를 지원한다.")

    add_callout(
        doc,
        "구현 원칙",
        "AI 결과는 원문을 대체하지 않는다. 사용자가 버튼을 눌렀을 때만 생성하며, 사실 확인은 원문과 출처를 기준으로 한다.",
    )

    # Page 4: functions part 2 and results
    page_break(doc)
    add_heading(doc, "2. 기능 구현 및 결과 (계속)", 1)
    add_heading(doc, "2.4 뉴스레터 구독과 정기 발송", 2)
    add_paragraph(
        doc,
        "사용자는 기본 또는 사용자 정의 카테고리를 선택해 이메일을 등록할 수 있다. 서버는 최근 RSS 항목을 수집하고, "
        "본문 추출과 AI 요약이 가능한 경우 이를 활용해 텍스트·HTML 뉴스레터를 구성한다. SMTP 설정이 없을 때는 "
        "구독 정보만 저장하고 안내 메시지를 반환한다. 정기 발송은 외부 cron이나 플랫폼 스케줄러에서 CLI 또는 "
        "관리자 API를 호출하는 방식으로 설계하였다.",
    )

    newsletter_rows = [
        ["구독", "이메일·카테고리·주간 주기 저장", "SQLite"],
        ["첫 메일", "구독 직후 발송 성공 시 현재 주차 완료 기록", "중복 방지"],
        ["주간 배치", "활성 weekly 구독 중 미처리 대상만 조회", "CLI/API"],
        ["발송 기록", "claimed·sent·failed 상태와 오류 저장", "SQLite"],
        ["해지", "토큰으로 구독 상태를 inactive로 변경", "API"],
    ]
    add_table(doc, ["단계", "처리 내용", "구현 방식"], newsletter_rows, [1700, 5460, 2200], font_size=9.0)

    add_heading(doc, "2.5 사용자 화면과 접근성", 2)
    add_paragraph(
        doc,
        "뉴스 읽기, URL 입력, 리더뷰 화면을 분리하여 사용 목적을 명확히 하였다. 760px 이하에서는 내비게이션, "
        "기사 목록, 읽기 도구, AI 패널과 모달을 한 열로 재배치한다. 주요 입력에는 label을 연결하고, 로딩에는 "
        "aria-live, 오류에는 role=alert, 모달에는 dialog와 aria-modal을 적용하였다. 다만 tab 키보드 패턴, "
        "포커스 트랩과 실제 스크린리더 검증은 추가 과제로 남아 있다.",
    )

    add_heading(doc, "2.6 테스트 및 결과", 2)
    result_rows = [
        ["백엔드 단위 테스트", "12건", "전체 통과"],
        ["Python 문법 검사", "주요 5개 모듈", "통과"],
        ["프론트엔드 빌드", "Vite production build", "통과"],
        ["외부 서비스 실연동", "Claude·DeepL·SMTP·Supabase", "환경변수 필요"],
    ]
    add_table(doc, ["검증 항목", "범위", "결과"], result_rows, [2600, 4100, 2660], font_size=9.3)

    add_paragraph(
        doc,
        "자동 테스트에서는 URL 보정과 잘못된 프로토콜 거부, AI 관찰 노트 JSON 파싱, 낮은 신호 처리, "
        "주간 발송 대상 선별, 같은 주차의 중복 발송 방지, dry-run의 무부작용, SMTP 실패 기록 등을 확인하였다.",
        after=4,
    )
    add_callout(
        doc,
        "결과 요약",
        "기사 탐색부터 본문 읽기, AI 보조 기능, 사용자 카테고리, 뉴스레터 구독과 주간 배치까지 하나의 서비스 흐름으로 연결하였다.",
    )

    # Page 5: integration/data + maintenance
    page_break(doc)
    add_heading(doc, "3. 서비스 연동 및 데이터 처리", 1)
    add_heading(doc, "3.1 외부 서비스 연동", 2)
    service_rows = [
        ["RSS/Atom", "기사 목록·제목·링크·발행일·요약 수집", "피드 실패 시 다른 소스 계속 처리"],
        ["Claude API", "AI 요약·관찰 노트·피드 탐색 힌트", "API 키 없으면 해당 기능만 비활성"],
        ["DeepL API", "본문 번역", "키·사용량 오류를 사용자 메시지로 변환"],
        ["SMTP", "텍스트·HTML 뉴스레터 발송", "설정 없으면 구독 저장 후 안내"],
        ["Supabase", "공개 todos 테이블 조회 검증", "핵심 데이터 이전은 미완료"],
    ]
    add_table(doc, ["서비스", "용도", "예외 처리"], service_rows, [1800, 3300, 4260], font_size=8.8)

    add_heading(doc, "3.2 데이터 저장과 처리 흐름", 2)
    add_paragraph(
        doc,
        "SQLite에는 뉴스레터 카테고리, RSS 출처, 구독 정보와 발송 시도를 분리해 저장한다. "
        "발송 시도 테이블은 구독 ID와 ISO 주차 조합을 고유하게 제한하여 같은 구독을 같은 주에 두 번 처리하지 않도록 한다. "
        "RSS 데이터는 URL 중복 제거와 키워드 필터를 거친 뒤 최신순으로 정렬한다. 비밀키와 SMTP 계정은 코드에 저장하지 않고 "
        "환경변수로 주입한다.",
    )
    add_code_block(
        doc,
        [
            "RSS 수집 → 중복 제거·키워드 필터 → 기사 목록",
            "기사 선택 → HTML 추출 → 실패 시 RSS 본문 → 리더뷰",
            "구독 조회 → 주차 선점 → digest 생성 → SMTP → 결과 기록",
        ],
    )

    add_heading(doc, "4. 유지보수 계획", 1)
    maintenance_rows = [
        ["우선", "Supabase 선택 초기화", "환경변수가 없어도 핵심 리더뷰가 실행되도록 수정"],
        ["우선", "구독 중복 방지", "이메일+카테고리 고유 제약과 사용자 해지 링크 추가"],
        ["단기", "접근성 검증", "키보드, 포커스, 색 대비, 스크린리더 테스트"],
        ["단기", "배포 자동화", "프론트·백엔드 배포 및 외부 스케줄러 연결"],
        ["중기", "사용자 기능", "Auth, 북마크, 읽기 기록, 설정 동기화"],
        ["중기", "운영 관측", "실패 재시도, 로그, 발송 결과 조회 화면"],
    ]
    add_table(doc, ["우선순위", "항목", "계획"], maintenance_rows, [1500, 2700, 5160], font_size=8.8)

    # Page 6: conclusion and reflection
    page_break(doc)
    add_heading(doc, "5. 결론 및 고찰", 1)
    add_heading(doc, "5.1 개발 성과", 2)
    add_paragraph(
        doc,
        "‘읽을게’는 URL 본문 추출 MVP에서 시작해 RSS 기사 탐색, 접근성 중심 리더뷰, AI 보조 기능, "
        "사용자 카테고리와 뉴스레터까지 확장되었다. 가장 의미 있는 결과는 서로 다른 출처의 콘텐츠를 동일한 읽기 화면으로 "
        "연결하고, 기사 탐색과 이메일 구독을 같은 카테고리 모델로 통합한 점이다. 또한 주간 발송 과정에 선점과 상태 기록을 "
        "도입하여 단순 데모보다 운영에 가까운 구조를 경험하였다.",
    )

    add_heading(doc, "5.2 개발 과정에서 배운 점", 2)
    add_bullet(doc, "외부 사이트의 본문 추출은 한 가지 라이브러리만으로 해결되지 않으므로 HTML·메타데이터·RSS fallback이 필요했다.")
    add_bullet(doc, "AI 기능은 결과의 화려함보다 원문과 해석을 구분하고 실패 시 핵심 기능을 유지하는 설계가 중요했다.")
    add_bullet(doc, "이메일 발송은 보내는 기능보다 중복 방지, 실패 기록, 재실행 가능성을 먼저 설계해야 안정적이었다.")
    add_bullet(doc, "외부 DB 연동은 클라이언트 연결만으로 완료되지 않으며 인증, 스키마, 권한 정책, 마이그레이션이 함께 필요하다.")
    add_bullet(doc, "접근성은 ARIA 속성을 추가하는 것에서 끝나지 않고 실제 키보드와 보조기기 테스트가 필요하다.")

    add_heading(doc, "5.3 한계와 개선 방향", 2)
    add_paragraph(
        doc,
        "현재 핵심 데이터는 SQLite에 저장되므로 여러 서버 인스턴스나 사용자별 데이터 분리에는 한계가 있다. "
        "Supabase는 공개 테이블 조회 실험만 완료되었으며, 사용자 계정·북마크·읽기 기록은 구현되지 않았다. "
        "정기 뉴스레터도 애플리케이션 내부 스케줄러가 아니라 외부 cron 호출이 필요하고, 실패 발송의 자동 재시도와 "
        "메일 본문의 해지 링크 연결이 남아 있다. 향후에는 기능 수를 늘리기보다 배포, 인증, 접근성 검증과 운영 로그를 "
        "우선 보강하는 것이 적절하다.",
    )

    add_heading(doc, "5.4 최종 평가", 2)
    add_callout(
        doc,
        "최종 판단",
        "핵심 목표인 ‘웹 콘텐츠를 읽기 쉽고 접근성 좋은 형태로 제공한다’는 달성하였다. "
        "뉴스레터와 AI 기능까지 확장했지만, 실서비스 완성 단계가 아니라 배포·인증·데이터 이전이 남은 기능형 MVP로 평가한다.",
    )

    add_heading(doc, "참고 및 제출 자료", 2)
    refs = [
        "프로젝트 소스코드: rss-news 전체 폴더 또는 GitHub 저장소 링크",
        "실행 문서: README.md",
        "개발 일지: docs/reports/second-report.md, third-report.md, fourth-report.md",
        "주요 테스트: server/test_parse_request.py, test_article_insight.py, test_newsletter_delivery.py",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("작성 기준일: 2026. 6. 23.")
    set_run_font(run, size=9, italic=True, color=MID_GRAY)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.core_properties.title = "읽을게 최종 프로젝트 보고서"
    document.core_properties.subject = "접근성 중심 RSS 뉴스 리더 및 뉴스레터 서비스"
    document.core_properties.author = ""
    document.core_properties.keywords = "React, FastAPI, RSS, 접근성, 뉴스레터"
    document.save(OUTPUT)
    print(OUTPUT)
